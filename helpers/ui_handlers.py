"""
UI tool handlers. Display handlers stage elements/actions in pending_ui.
Ask handlers block until the user responds and return the result.
"""
import os
import chainlit as cl

from .session import stage_actions, stage_element
from .ui_builders import build_action, build_task


# --- Spreadsheet generation ---

async def generate_xlsx(args: dict) -> dict:
    import tempfile
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment

    filename = (args.get("filename") or "spreadsheet").removesuffix(".xlsx") + ".xlsx"
    sheets = args.get("sheets", [])
    if not sheets:
        sheets = [{"name": "Sheet1", "headers": [], "rows": []}]

    wb = Workbook()
    wb.remove(wb.active)

    for sheet_data in sheets:
        ws = wb.create_sheet(title=sheet_data.get("name") or "Sheet")
        headers = sheet_data.get("headers", [])
        rows = sheet_data.get("rows", [])

        if headers:
            ws.append(headers)
            header_fill = PatternFill("solid", fgColor="E2E8F0")
            for cell in ws[1]:
                cell.font = Font(bold=True)
                cell.fill = header_fill
                cell.alignment = Alignment(wrap_text=False)

        for row in rows:
            ws.append(list(row) if row else [])

        # Auto-width columns
        for col in ws.columns:
            max_len = max((len(str(c.value or "")) for c in col), default=0)
            ws.column_dimensions[col[0].column_letter].width = min(max_len + 4, 40)

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx", prefix="xlsx_")
    wb.save(tmp.name)
    tmp.close()

    safe_key = filename[:20].replace(" ", "_").lower()
    session_key = f"xlsx_{safe_key}"
    file_info = {"path": tmp.name, "filename": filename}
    cl.user_session.set(session_key, file_info)
    cl.user_session.set("latest_xlsx", file_info)

    element = cl.CustomElement(
        name="XlsxPreview",
        props={"filename": filename, "sheets": sheets, "session_key": session_key},
        display="inline",
    )
    stage_element(element)
    total_rows = sum(len(s.get("rows", [])) for s in sheets)
    return {"success": True, "preview_displayed": True, "filename": filename, "sheet_count": len(sheets), "total_rows": total_rows}


# --- Utility ---

async def rename_thread(args: dict) -> dict:
    name = args.get("name", "").strip()
    if not name:
        return {"success": False, "error": "name is required"}
    thread_id = cl.context.session.thread_id
    await cl.context.emitter.emit("first_interaction", {"interaction": name, "thread_id": thread_id})
    return {"success": True, "name": name}


# --- Dynamic custom element code generation ---

_GENERATED_ELEMENT_PREFIX = "_tmp_"


def _jsx_hoisting_error(jsx_code: str) -> str | None:
    import re

    for m in re.finditer(r"useState\s*\(\s*(?:\(\)\s*=>\s*)?(\w+)\s*\(", jsx_code):
        fn = m.group(1)
        if fn in ("useState", "useEffect", "useRef", "useMemo", "useCallback"):
            continue
        use_pos = m.start()
        before = jsx_code[:use_pos]
        if re.search(rf"function\s+{re.escape(fn)}\s*\(", before):
            continue
        if re.search(rf"const\s+{re.escape(fn)}\s*=", before):
            continue
        if re.search(rf"const\s+{re.escape(fn)}\s*=", jsx_code[use_pos:]):
            return (
                f"useState calls {fn}() before {fn} is defined (causes runtime TDZ error). "
                f"Move helpers outside the component, use function declarations, "
                f"or use useState(() => ...) only after function declarations above useState."
            )
    return None


async def generate_custom_element_code(args: dict) -> dict:
    import re
    import uuid

    element_name = args.get("element_name", "").strip()
    jsx_code     = args.get("jsx_code", "").strip()
    props        = args.get("props", {})

    if not element_name or not re.match(r"^[A-Z][a-zA-Z0-9]*$", element_name):
        return {"error": f"element_name must be PascalCase (e.g. MyWidget), got: {element_name!r}"}
    if re.search(r"tetris", element_name, re.I):
        return {
            "error": "Use display_custom_element with name='TetrisGame'. Do not generate Tetris JSX.",
        }

    _grid_block_game = re.search(
        r"tetris|tetromino|grid\s*game|puzzle\s*grid|breakout|snake|minesweeper|sudoku|cell\s*===?\s*1|board\s*\[",
        f"{element_name} {jsx_code}",
        re.I,
    )
    _themed_sprite = re.search(
        r"dinosaur|dragon|cliff|cactus|obstacle|rex|flappy|mario|sprite|character|enemy",
        f"{element_name} {jsx_code}",
        re.I,
    )
    _has_icon_art = re.search(
        r"<img\s|api\.iconify\.design|from\s+['\"]lucide-react['\"]",
        jsx_code,
        re.I,
    )
    if _themed_sprite and not _grid_block_game and not _has_icon_art:
        return {
            "error": (
                "This UI depicts named subjects (dinosaur, cactus, etc.) but uses no icons. "
                "Call fetch_iconify_icons first, then embed svg_url in <img> tags. "
                "Grid games (Tetris, Snake) should use colored cells — not Iconify."
            ),
        }
    if re.search(r"// Logic to|TODO|FIXME|not implemented", jsx_code, re.I):
        return {
            "error": (
                "JSX contains stub or placeholder logic. Implement the full UI — "
                "no empty functions or '// Logic to...' comments."
            ),
            "required_fix": "Implement all game/UI behavior (movement, collision, scoring, etc.) with real code.",
        }
    if re.search(r"function\s+\w+\s*\([^)]*\)\s*\{\s*\}", jsx_code):
        return {"error": "Empty function bodies are not allowed — implement complete logic."}

    hoisting_err = _jsx_hoisting_error(jsx_code)
    if hoisting_err:
        return {"error": hoisting_err}
    if re.search(r'window\.addEventListener\s*\(\s*["\']keydown', jsx_code):
        if not re.search(r"onBlur|onFocus|activeRef|isContentEditable", jsx_code, re.I):
            return {
                "error": (
                    "Global keydown steals Space/arrows from chat. "
                    "Gate keys with onFocus/onBlur + activeRef on the root Card (see tool description)."
                ),
            }
    if not jsx_code:
        return {"error": "jsx_code is required"}

    _ALLOWED_UI = (
        "button", "card", "badge", "input", "textarea", "label", "separator", "checkbox",
        "select", "progress", "tabs", "switch", "avatar", "skeleton", "tooltip",
        "scroll-area", "table", "dialog", "dropdown-menu",
    )
    for m in re.finditer(r"@/components/ui/([\w-]+)", jsx_code):
        if m.group(1).lower() not in _ALLOWED_UI:
            return {
                "error": (
                    f"Invalid @/components/ui/{m.group(1)} import. "
                    f"Allowed: {', '.join(_ALLOWED_UI)}. "
                    "Use lucide-react for icons, never emojis."
                ),
            }

    invalid_lucide = {
        "Pumpkin", "Bat", "Spider", "Halloween", "CandyCorn", "Witch", "Zombie",
        "Christmas", "Santa", "Reindeer", "Tree",
    }
    for m in re.finditer(r"import\s*\{([^}]+)\}\s*from\s*['\"]lucide-react['\"]", jsx_code):
        for name in re.split(r",\s*", m.group(1)):
            name = name.strip().split(" as ")[0].strip()
            if name in invalid_lucide:
                return {
                    "error": (
                        f"Lucide icon '{name}' does not exist (causes React error #130). "
                        "Use Ghost, Skull, Moon, Star, Flame, Sparkles, Skull, Heart, etc."
                    ),
                }

    if re.search(r"[\U0001F300-\U0001FAFF\u2600-\u27BF]", jsx_code):
        return {"error": "Emojis are not allowed in generated JSX. Use Lucide icons instead."}

    if re.search(r"Decoration\s+\d+", jsx_code, re.I):
        return {
            "error": (
                "Do not use placeholder labels like 'Decoration 1'. "
                "Build real visuals (shapes, positioned ornaments, icons) that match the user's request."
            ),
        }

    if re.search(r"(what is|solve|answer).*\d+\s*[\+\-\*]\s*\d+|captcha.*<Input|verify.*math", jsx_code, re.I):
        return {
            "error": (
                "Do not build a math-quiz captcha. "
                "Use a visual selection challenge (e.g. tile grid with Lucide icons and Verify)."
            ),
        }

    elements_dir = os.path.join(
        os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
        "public", "elements",
    )
    os.makedirs(elements_dir, exist_ok=True)
    base_stem = f"{_GENERATED_ELEMENT_PREFIX}{element_name}"
    unique_id = uuid.uuid4().hex[:8]
    file_stem = f"{base_stem}_{unique_id}"
    file_path = os.path.join(elements_dir, f"{file_stem}.jsx")

    if base_stem in jsx_code:
        jsx_code = jsx_code.replace(base_stem, file_stem)

    with open(file_path, "w", encoding="utf-8") as f:
        f.write(jsx_code)

    element = cl.CustomElement(name=file_stem, props=props, display="inline")
    stage_element(element)
    return {
        "success": True,
        "element_name": file_stem,
        "file_written": f"public/elements/{file_stem}.jsx",
        "preview_displayed": True,
        "do_not_call": "display_custom_element",
    }


# --- Custom Plotly chart from LLM-provided data ---

async def display_plotly_chart(args: dict) -> dict:
    import plotly.graph_objects as go

    title       = args.get("title", "Chart")
    chart_type  = args.get("chart_type", "bar")
    x_labels    = args.get("x_labels", [])
    y_values    = args.get("y_values", [])
    x_label     = args.get("x_axis_label", "")
    y_label     = args.get("y_axis_label", "")
    color       = args.get("color", "#6366f1")

    if chart_type == "line":
        trace = go.Scatter(x=x_labels, y=y_values, mode="lines+markers",
                           line=dict(color=color, width=2),
                           marker=dict(size=7, color=color))
    elif chart_type == "pie":
        trace = go.Pie(labels=x_labels, values=y_values)
    else:
        trace = go.Bar(x=x_labels, y=y_values, marker_color=color)

    fig = go.Figure(trace)
    fig.update_layout(
        title=dict(text=title, font=dict(size=15)),
        xaxis_title=x_label,
        yaxis_title=y_label,
        template="plotly_dark",
        margin=dict(l=48, r=24, t=52, b=40),
        height=340,
    )

    element = cl.Plotly(name=title, figure=fig, display="inline")
    stage_element(element)
    return {"success": True, "points": len(x_labels)}


# --- Visual timeline (Gantt) ---

async def display_gantt_timeline(args: dict) -> dict:
    element = cl.CustomElement(name="GanttTimeline", props=args, display="inline")
    stage_element(element)
    return {"success": True, "item_count": len(args.get("items", []))}


# --- Document generation ---

async def generate_word_doc(args: dict) -> dict:
    import tempfile
    from docx import Document

    title = args.get("title", "Document")
    filename = (args.get("filename") or "document").rstrip(".docx") + ".docx"
    sections = args.get("sections", [])

    doc = Document()
    doc.add_heading(title, level=0)

    for section in sections:
        heading = section.get("heading", "")
        if heading:
            doc.add_heading(heading, level=1)
        for block in section.get("content", []):
            btype = block.get("type")
            if btype == "paragraph":
                doc.add_paragraph(block.get("text", ""))
            elif btype == "bullets":
                for item in block.get("items", []):
                    doc.add_paragraph(item, style="List Bullet")
            elif btype == "table":
                headers = block.get("headers", [])
                rows = block.get("rows", [])
                if headers:
                    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
                    table.style = "Table Grid"
                    for i, h in enumerate(headers):
                        table.rows[0].cells[i].text = h
                    for ri, row in enumerate(rows):
                        for ci, val in enumerate(row):
                            if ci < len(table.rows[ri + 1].cells):
                                table.rows[ri + 1].cells[ci].text = str(val)

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx", prefix="worddoc_")
    doc.save(tmp.name)
    tmp.close()

    safe_key = title[:20].replace(" ", "_").lower()
    session_key = f"word_doc_{safe_key}"
    doc_info = {"path": tmp.name, "filename": filename, "title": title}
    cl.user_session.set(session_key, doc_info)
    cl.user_session.set("latest_word_doc", doc_info)

    element = cl.CustomElement(
        name="WordDocPreview",
        props={"title": title, "filename": filename, "sections": sections, "session_key": session_key},
        display="inline",
    )
    stage_element(element)
    return {"success": True, "preview_displayed": True, "filename": filename, "section_count": len(sections)}


# --- Display tools ---

async def display_actions(args: dict) -> dict:
    buttons = args.get("buttons", [])
    if not buttons or len(buttons) > 5:
        return {"success": False, "error": "Need 1-5 buttons"}
    stage_actions(buttons)
    return {"success": True, "button_count": len(buttons)}


async def display_dataframe(args: dict) -> dict:
    import pandas as pd

    key = args.get("session_key", "latest_dataframe")
    df = cl.user_session.get(key)
    if df is None:
        return {"success": False, "error": f"No dataframe found under session key '{key}'. Call fetch_sample_dataframe first."}

    name = args.get("name", "table")
    display = args.get("display", "inline")
    element = cl.Dataframe(name=name, data=df, display=display)
    stage_element(element)
    return {"success": True, "name": name, "rows": len(df)}


async def display_plotly(args: dict) -> dict:
    import plotly

    key = args.get("session_key", "latest_plotly_figure")
    fig_json = cl.user_session.get(key)
    if fig_json is None:
        return {"success": False, "error": f"No figure found under session key '{key}'. Call build_sample_plotly_figure first."}

    name = args.get("name", "chart")
    display = args.get("display", "inline")
    fig = plotly.io.from_json(fig_json)
    element = cl.Plotly(name=name, figure=fig, display=display)
    stage_element(element)
    return {"success": True, "name": name}


def _style_pyplot_figure(fig, ax, chart_type: str, title: str, x_label: str, y_label: str):
    accent = "#6366f1"
    muted = "#94a3b8"
    grid = "#e2e8f0"
    text = "#0f172a"

    fig.patch.set_facecolor("#ffffff")
    ax.set_facecolor("#ffffff")
    for side in ("top", "right"):
        ax.spines[side].set_visible(False)
    ax.spines["left"].set_color(grid)
    ax.spines["bottom"].set_color(grid)
    ax.tick_params(colors=muted, labelsize=9, length=0, pad=6)
    if chart_type != "pie":
        ax.yaxis.grid(True, linestyle="-", linewidth=0.8, color=grid, alpha=0.9)
        ax.set_axisbelow(True)
        if x_label:
            ax.set_xlabel(x_label, color=muted, fontsize=10, labelpad=8)
        if y_label:
            ax.set_ylabel(y_label, color=muted, fontsize=10, labelpad=8)
    if title:
        ax.set_title(title, fontsize=13, fontweight=600, color=text, pad=12)
    return accent


async def display_pyplot(args: dict) -> dict:
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    chart_type = args.get("chart_type", "bar")
    title = args.get("title", "")
    x_labels = args.get("x_labels", ["A", "B", "C", "D"])
    y_values = args.get("y_values", [10, 25, 15, 30])
    x_label = args.get("x_label", "")
    y_label = args.get("y_label", "")
    name = args.get("name", "pyplot")
    display = args.get("display", "inline")

    fig, ax = plt.subplots(figsize=(7, 3.75), dpi=120)
    accent = _style_pyplot_figure(fig, ax, chart_type, title, x_label, y_label)
    x_pos = list(range(len(x_labels)))

    if chart_type == "pie":
        colors = ["#6366f1", "#818cf8", "#a5b4fc", "#c7d2fe", "#e0e7ff", "#eef2ff"]
        wedges, _, autotexts = ax.pie(
            y_values,
            labels=x_labels,
            autopct="%1.0f%%",
            colors=colors[: len(y_values)],
            startangle=90,
            pctdistance=0.75,
            wedgeprops={"linewidth": 1, "edgecolor": "#ffffff"},
        )
        for t in autotexts:
            t.set_fontsize(9)
            t.set_color("#334155")
        ax.axis("equal")
    elif chart_type == "line":
        ax.plot(
            x_pos,
            y_values,
            color=accent,
            linewidth=2.25,
            marker="o",
            markersize=7,
            markerfacecolor="#ffffff",
            markeredgewidth=2,
            markeredgecolor=accent,
            zorder=3,
        )
        ax.set_xticks(x_pos)
        ax.set_xticklabels(x_labels)
        if len(x_labels) > 5:
            plt.setp(ax.get_xticklabels(), rotation=30, ha="right")
        ax.margins(x=0.05)
    elif chart_type == "scatter":
        ax.scatter(x_pos, y_values, color=accent, s=56, alpha=0.85, edgecolors="#ffffff", linewidths=1.5)
        ax.set_xticks(x_pos)
        ax.set_xticklabels(x_labels)
    else:  # bar
        ax.bar(
            x_pos,
            y_values,
            color=accent,
            width=0.55,
            edgecolor="none",
            alpha=0.92,
            zorder=3,
        )
        ax.set_xticks(x_pos)
        ax.set_xticklabels(x_labels)
        if len(x_labels) > 5:
            plt.setp(ax.get_xticklabels(), rotation=30, ha="right")

    plt.tight_layout(pad=1.2)
    element = cl.Pyplot(name=name, figure=fig, display=display)
    stage_element(element)
    plt.close(fig)
    return {"success": True, "name": name, "chart_type": chart_type}


async def display_pdf(args: dict) -> dict:
    url = args.get("url", "https://www.w3.org/WAI/WCAG21/wcag21.pdf")
    name = args.get("name", "document.pdf")
    display = args.get("display", "inline")
    element = cl.Pdf(name=name, url=url, display=display)
    stage_element(element)
    return {"success": True, "name": name, "url": url}


async def display_file(args: dict) -> dict:
    key = args.get("session_key", "latest_file_path")
    name = args.get("name")
    text_content = args.get("text_content")

    import mimetypes

    def _mime(filename: str) -> str:
        mime, _ = mimetypes.guess_type(filename)
        return mime or "application/octet-stream"

    if text_content:
        fname = name or "output.txt"
        element = cl.File(
            name=fname,
            content=text_content.encode("utf-8"),
            mime=_mime(fname),
            display="inline",
        )
        stage_element(element)
        return {"success": True, "name": fname, "source": "inline"}

    path = cl.user_session.get(key)
    if not path or not os.path.exists(path):
        return {"success": False, "error": f"No file path under session key '{key}'. Call get_sample_file_path first."}

    filename = name or os.path.basename(path)
    element = cl.File(name=filename, path=path, mime=_mime(filename), display="inline")
    stage_element(element)
    return {"success": True, "name": filename, "path": path}


def _element_jsx_path(name: str) -> str:
    base = os.path.join(
        os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
        "public", "elements",
    )
    return os.path.join(base, f"{name}.jsx")


# Block curated runners/captcha from display when user wants generated variety — except TetrisGame (fully working built-in).
_BUILT_IN_GAME_ELEMENTS = frozenset({"PlatformGame", "TileCaptcha"})


async def display_custom_element(args: dict) -> dict:
    name = args["name"]
    props = args.get("props", {})
    display = args.get("display", "inline")
    if name in _BUILT_IN_GAME_ELEMENTS and not args.get("demo"):
        return {
            "error": (
                f"Built-in {name} is disabled for open-ended requests. "
                "Use generate_custom_element_code with mechanics and visuals specific to what the user asked."
            ),
        }
    jsx_path = _element_jsx_path(name)
    if not os.path.isfile(jsx_path):
        return {
            "error": (
                f"No JSX at public/elements/{name}.jsx (404 in UI). "
                "Use generate_custom_element_code or an existing built-in."
            ),
        }
    element = cl.CustomElement(name=name, props=props, display=display)
    # Store reference so Python can update props later
    cl.user_session.set(f"custom_element_{name}", element)
    stage_element(element)
    return {"success": True, "element": name, "props": list(props.keys())}


async def display_tasklist(args: dict) -> dict:
    tasks_raw = args.get("tasks", [])
    title = args.get("title", "Tasks")

    task_list = cl.TaskList()
    task_list.status = title

    for t in tasks_raw:
        await task_list.add_task(build_task(t))

    await task_list.send()
    return {"success": True, "task_count": len(tasks_raw), "status": title}


async def display_text_element(args: dict) -> dict:
    name = args["name"]
    content = args["content"]
    display = args.get("display", "side")
    language = args.get("language")
    element = cl.Text(name=name, content=content, display=display, language=language)
    stage_element(element)
    return {"success": True, "name": name}


async def display_code_block(args: dict) -> dict:
    code = args.get("code") or args.get("content")
    if not code or not str(code).strip():
        return {"error": "code (or content) is required."}
    element = cl.CustomElement(name="CodeBlock", props=args, display="inline")
    stage_element(element)
    line_count = len(str(code).splitlines())
    return {
        "success": True,
        "line_count": line_count,
        "collapsible": line_count > (args.get("max_lines") or args.get("maxLines") or 12),
    }


async def display_dropdown(args: dict) -> dict:
    options = args.get("options", [])
    element = cl.CustomElement(
        name="DropdownSelect",
        props={
            "options": options,
            "title": args.get("title", "Select an option"),
            "placeholder": args.get("placeholder", "Choose..."),
        },
        display="inline",
    )
    stage_element(element)
    return {"success": True, "option_count": len(options)}


async def display_checkbox(args: dict) -> dict:
    options = args.get("options", [])
    element = cl.CustomElement(
        name="CheckboxGroup",
        props={
            "options": options,
            "title": args.get("title", "Select options"),
            "multiple": args.get("multiple", True),
            "min_selected": args.get("min_selected", 0),
            "max_selected": args.get("max_selected"),
        },
        display="inline",
    )
    stage_element(element)
    return {"success": True, "option_count": len(options)}


async def display_rank(args: dict) -> dict:
    items = args.get("items", [])
    element = cl.CustomElement(
        name="DraggableRank",
        props={
            "items": items,
            "title": args.get("title", "Rank these items"),
            "instruction": args.get("instruction", "Drag to reorder. Top = highest priority."),
        },
        display="inline",
    )
    stage_element(element)
    return {"success": True, "item_count": len(items)}


async def display_categories(args: dict) -> dict:
    columns = args.get("columns", [])
    if len(columns) < 2 or len(columns) > 3:
        return {"success": False, "error": "Need 2 or 3 columns"}
    element = cl.CustomElement(
        name="DraggableCategories",
        props={
            "columns": columns,
            "title": args.get("title", "Categorize these items"),
            "instruction": args.get("instruction", "Drag pills between columns to categorize them."),
        },
        display="inline",
    )
    stage_element(element)
    total = sum(len(c.get("items", [])) for c in columns)
    return {"success": True, "column_count": len(columns), "total_items": total}

def _normalize_record_editor_props(args: dict) -> dict:
    props = dict(args)
    if props.get("key_results") and not props.get("items"):
        props["items"] = props.pop("key_results")
    if props.get("key_results"):
        props.pop("key_results", None)
    if not props.get("items_label") and (props.get("items") or props.get("icon") == "target" or "goal" in (props.get("title") or "").lower()):
        props.setdefault("items_label", "Key results")
    if not props.get("icon") and props.get("items_label") == "Key results":
        props.setdefault("icon", "target")
    return props


async def display_record_editor(args: dict) -> dict:
    props = _normalize_record_editor_props(args)
    if not props.get("title"):
        return {"error": "title is required."}
    element = cl.CustomElement(name="RecordEditor", props=props, display="inline")
    cl.user_session.set("custom_element_RecordEditor", element)
    stage_element(element)
    return {"success": True, "title": props.get("title")}


async def display_goal_card(args: dict) -> dict:
    return await display_record_editor(args)


async def display_stat_grid(args: dict) -> dict:
    stats = args.get("stats") or args.get("metrics") or []
    if not stats:
        return {"error": "stats (or metrics) is required — at least one metric."}
    element = cl.CustomElement(name="StatGrid", props=args, display="inline")
    stage_element(element)
    return {"success": True, "stat_count": len(stats)}


async def display_progress_list(args: dict) -> dict:
    props = dict(args)
    items = props.get("items") or props.get("goals") or []
    if not items:
        return {"error": "items (or goals) is required — at least one row to show."}
    element = cl.CustomElement(name="ProgressList", props=props, display="inline")
    stage_element(element)
    return {"success": True, "item_count": len(items)}


async def display_goal_progress(args: dict) -> dict:
    return await display_progress_list(args)


def _normalize_highlight_props(args: dict) -> dict:
    props = dict(args)
    if props.get("award_name") and not props.get("headline"):
        props["headline"] = props["award_name"]
    if props.get("citation") and not props.get("body"):
        props["body"] = props["citation"]
    if (props.get("show_nominate") or props.get("show_cta")) and not props.get("cta"):
        props["cta"] = {
            "label": "Nominate someone",
            "fields": [
                {"id": "name", "label": "Name", "type": "text"},
                {"id": "reason", "label": "Reason", "type": "textarea"},
            ],
            "submit_label": "Submit nomination",
        }
    return props


async def display_highlight_card(args: dict) -> dict:
    props = _normalize_highlight_props(args)
    if not (props.get("headline") or props.get("body")):
        return {"error": "headline or body is required."}
    element = cl.CustomElement(name="HighlightCard", props=props, display="inline")
    stage_element(element)
    return {"success": True, "headline": props.get("headline")}


async def display_tree(args: dict) -> dict:
    props = dict(args)
    nodes = props.get("nodes") or []
    if not nodes:
        return {"error": "nodes is required — at least one root node."}
    element = cl.CustomElement(name="NestedTree", props=props, display="inline")
    stage_element(element)
    return {"success": True, "node_count": len(nodes)}


async def display_okr_tree(args: dict) -> dict:
    return await display_tree(args)


# display_feedback_picker archived — use display_card_picker


async def display_card_picker(args: dict) -> dict:
    element = cl.CustomElement(name="CardPicker", props=args, display="inline")
    stage_element(element)
    return {"success": True, "item_count": len(args.get("items", []))}


def _normalize_rating_matrix_props(args: dict) -> dict:
    props = dict(args)
    if not props.get("items") and props.get("competencies"):
        props["items"] = [
            {"id": c["id"], "label": c.get("name") or c.get("label", c["id"]), "description": c.get("description")}
            for c in props["competencies"]
        ]
    props.pop("competencies", None)
    return props


async def display_rating_matrix(args: dict) -> dict:
    props = _normalize_rating_matrix_props(args)
    items = props.get("items") or []
    if not items:
        return {"error": "items is required — at least one row to rate."}
    element = cl.CustomElement(name="RatingMatrix", props=props, display="inline")
    stage_element(element)
    return {"success": True, "item_count": len(items)}


async def display_competency_rater(args: dict) -> dict:
    return await display_rating_matrix(args)


# display_feedback_card archived — use display_section_card
# display_feedback_summary archived — use display_pill_board


async def display_section_card(args: dict) -> dict:
    element = cl.CustomElement(name="SectionCard", props=args, display="inline")
    stage_element(element)
    return {"success": True, "section_count": len(args.get("sections", []))}


async def display_pill_board(args: dict) -> dict:
    element = cl.CustomElement(name="PillBoard", props=args, display="inline")
    stage_element(element)
    return {"success": True, "column_count": len(args.get("columns", []))}


def _normalize_form_props(args: dict) -> dict:
    props = dict(args)
    fields = list(props.get("fields") or [])
    if not fields and props.get("sections"):
        fields = [
            {
                "id": s["id"],
                "label": s.get("title") or s.get("label", s["id"]),
                "placeholder": s.get("placeholder", ""),
                "type": s.get("type", "textarea"),
                "required": s.get("required", True),
                "value": s.get("value", ""),
                "max_chars": s.get("max_chars"),
                "options": s.get("options"),
                "rows": s.get("rows"),
            }
            for s in props["sections"]
        ]
    seen = {f.get("id") for f in fields if f.get("id")}
    for row in props.get("rows") or []:
        for ref in row.get("fields") or row.get("columns") or []:
            if isinstance(ref, dict) and ref.get("id") and ref["id"] not in seen:
                fields.append(ref)
                seen.add(ref["id"])
    props["fields"] = fields
    if props.get("employee_name") and not props.get("subtitle"):
        props["subtitle"] = props["employee_name"]
    props.pop("sections", None)
    props.pop("employee_name", None)
    return props


async def display_form(args: dict) -> dict:
    props = _normalize_form_props(args)
    fields = props.get("fields") or []
    has_rows = bool(props.get("rows"))
    if not fields and not has_rows:
        return {"error": "fields, rows, or sections is required — provide at least one field for the user to fill in."}
    element = cl.CustomElement(name="DynamicForm", props=props, display="inline")
    cl.user_session.set("custom_element_DynamicForm", element)
    stage_element(element)
    return {"success": True, "field_count": len(fields), "title": props.get("title")}


async def display_assessment_form(args: dict) -> dict:
    return await display_form(args)


def _normalize_tier_picker_props(args: dict) -> dict:
    props = dict(args)
    if props.get("employee") and not props.get("subtitle"):
        parts = [props.pop("employee", None), props.pop("cycle", None)]
        props["subtitle"] = " - ".join(p for p in parts if p)
    if "require_justification" in props and "require_comment" not in props:
        props["require_comment"] = props["require_justification"]
    if props.get("current_rating") and not props.get("current_value"):
        props["current_value"] = props["current_rating"]
    return props


async def display_tier_picker(args: dict) -> dict:
    props = _normalize_tier_picker_props(args)
    if not props.get("tiers"):
        return {"error": "tiers is required — provide label, value, and optional description per option."}
    element = cl.CustomElement(name="TierPicker", props=props, display="inline")
    stage_element(element)
    return {"success": True, "tier_count": len(props["tiers"])}


async def display_performance_rating(args: dict) -> dict:
    return await display_tier_picker(args)


def _normalize_editable_list_props(args: dict) -> dict:
    props = dict(args)
    if props.get("topics") and not props.get("rows"):
        props["preset"] = props.get("preset") or "agenda"
    elif props.get("items") and not props.get("rows"):
        props["preset"] = props.get("preset") or "tasks"
    if props.get("employee") and not props.get("subtitle"):
        props["subtitle"] = props.get("employee")
    return props


async def display_editable_list(args: dict) -> dict:
    props = _normalize_editable_list_props(args)
    rows = props.get("rows") or props.get("topics") or props.get("items") or []
    if not rows and props.get("preset") not in ("agenda", "tasks"):
        pass  # empty list ok — user can add rows
    element = cl.CustomElement(name="EditableList", props=props, display="inline")
    stage_element(element)
    n = len(rows)
    return {"success": True, "row_count": n, "preset": props.get("preset")}


async def display_checkin_agenda(args: dict) -> dict:
    args = dict(args)
    args.setdefault("preset", "agenda")
    return await display_editable_list(args)


async def display_action_items(args: dict) -> dict:
    args = dict(args)
    args.setdefault("preset", "tasks")
    return await display_editable_list(args)


async def display_award_card(args: dict) -> dict:
    return await display_highlight_card(args)


async def display_timeline(args: dict) -> dict:
    props = dict(args)
    if props.get("employee") and not props.get("subtitle"):
        props["subtitle"] = props.pop("employee")
    element = cl.CustomElement(name="Timeline", props=props, display="inline")
    stage_element(element)
    return {"success": True, "event_count": len(props.get("events", []))}


async def display_date_list(args: dict) -> dict:
    props = dict(args)
    if not props.get("items") and props.get("deadlines"):
        props["items"] = props["deadlines"]
    element = cl.CustomElement(name="DateList", props=props, display="inline")
    stage_element(element)
    return {"success": True, "item_count": len(props.get("items") or [])}


async def display_deadline_countdown(args: dict) -> dict:
    return await display_date_list(args)


async def display_year_calendar(args: dict) -> dict:
    import datetime
    year = args.get("year") or datetime.date.today().year
    title = args.get("title", f"{year} Calendar")
    element = cl.CustomElement(
        name="YearCalendar",
        props={"year": year, "title": title},
        display="inline",
    )
    cl.user_session.set("custom_element_YearCalendar", element)
    stage_element(element)
    return {"success": True, "year": year}


# --- Ask tools ---

async def ask_user(args: dict) -> dict:
    content = args["content"]
    timeout = args.get("timeout", 90)
    res = await cl.AskUserMessage(content=content, timeout=timeout).send()
    if res is None:
        return {"timed_out": True}
    return {"response": res.get("output") or res.get("content", ""), "timed_out": False}


async def ask_file(args: dict) -> dict:
    content = args["content"]
    accept = args.get("accept", ["text/plain", "application/pdf", "image/*"])
    max_files = args.get("max_files", 1)
    timeout = args.get("timeout", 90)

    files = await cl.AskFileMessage(
        content=content,
        accept=accept,
        max_files=max_files,
        timeout=timeout,
    ).send()

    if files is None:
        return {"timed_out": True, "files": []}
    return {
        "timed_out": False,
        "files": [{"name": f.name, "path": f.path, "size": f.size, "type": f.type} for f in files],
    }


async def ask_action(args: dict) -> dict:
    content = args["content"]
    timeout = args.get("timeout", 90)
    actions = [
        cl.Action(
            name=a["name"],
            label=a["label"],
            icon=a.get("icon"),
            payload=a.get("payload", {}),
        )
        for a in args.get("actions", [])
    ]
    res = await cl.AskActionMessage(content=content, actions=actions, timeout=timeout).send()
    if res is None:
        return {"timed_out": True}
    return {"timed_out": False, "name": res.get("name"), "payload": res.get("payload", {})}


async def ask_element(args: dict) -> dict:
    content = args["content"]
    timeout = args.get("timeout", 90)
    title = args.get("title", "Please fill out this form")
    fields = args.get("fields") or [
        {"id": "response", "label": "Your response", "type": "text", "required": True}
    ]

    element = cl.CustomElement(
        name="DynamicForm",
        display="inline",
        props={"title": title, "fields": fields, "timeout": timeout, "ask_mode": True, "compact": True},
    )
    res = await cl.AskElementMessage(content=content, element=element, timeout=timeout).send()
    if res is None:
        return {"timed_out": True, "submitted": False}
    return {"timed_out": False, "submitted": res.get("submitted", False), **{k: v for k, v in res.items() if k != "submitted"}}
